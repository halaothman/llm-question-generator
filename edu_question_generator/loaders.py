"""Extract text from PDF, DOCX, and TXT files via pdfplumber/python-docx."""
from __future__ import annotations

import os

import docx
import pdfplumber


def _format_table(table: list[list | None]) -> str:
    """Convert an extracted table to markdown-like text."""
    rows: list[list[str]] = []
    for row in table:
        if not row:
            continue
        rows.append([str(cell or "").replace("\n", " ").strip() for cell in row])
    if not rows:
        return ""
    lines = ["| " + " | ".join(row) + " |" for row in rows]
    return "\n".join(lines)


def load_pdf(path: str) -> str:
    """Read PDF: page text + tables (no OCR)."""
    parts: list[str] = []

    with pdfplumber.open(path) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            parts.append(f"\n--- صفحة {page_num} ---\n")

            text = (page.extract_text() or "").strip()
            if text:
                parts.append(text)

            for table_index, table in enumerate(page.extract_tables() or [], 1):
                formatted = _format_table(table)
                if formatted.strip():
                    parts.append(f"\n[جدول {table_index}]\n{formatted}\n")

    return "\n".join(parts).strip()


def load_docx(path: str) -> str:
    """Read DOCX: paragraphs + tables."""
    document = docx.Document(path)
    parts: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())

    for table_index, table in enumerate(document.tables, 1):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        formatted = _format_table(rows)
        if formatted.strip():
            parts.append(f"\n[جدول {table_index}]\n{formatted}\n")

    return "\n".join(parts).strip()


def load_text(path: str) -> str:
    """Load a document: dispatches to load_pdf, load_docx, or UTF-8 by extension.

    Raises:
        ValueError: If the file extension is not supported.
    """
    ext = os.path.splitext(path)[1].lower()

    if ext == ".pdf":
        return load_pdf(path)

    if ext in {".docx", ".doc"}:
        return load_docx(path)

    if ext in {".txt", ".md"}:
        with open(path, encoding="utf-8", errors="ignore") as f:
            return f.read()

    raise ValueError(f"Unsupported file type: {ext}")
